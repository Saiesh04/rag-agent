"""
Document processing utilities for handling PDF, DOCX, and Excel files
"""
import os
import tempfile
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
import pandas as pd
import PyPDF2
from docx import Document
import hashlib
from datetime import datetime

class DocumentProcessor:
    """Handles document processing for various file types"""
    
    def __init__(self):
        self.supported_extensions = {'.pdf', '.docx', '.xlsx', '.xls', '.txt'}
    
    def is_supported(self, file_path: str) -> bool:
        """Check if file type is supported"""
        return Path(file_path).suffix.lower() in self.supported_extensions
    
    def extract_text_from_pdf(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from PDF file"""
        text = ""
        metadata = {"pages": 0, "file_type": "pdf"}
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                metadata["pages"] = len(pdf_reader.pages)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                    
        except Exception as e:
            raise Exception(f"Error processing PDF: {str(e)}")
        
        return text.strip(), metadata
    
    def extract_text_from_docx(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text_parts = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text.strip())
            
            text = "\n\n".join(text_parts)
            metadata = {
                "paragraphs": len(doc.paragraphs),
                "file_type": "docx"
            }
            
            return text, metadata
            
        except Exception as e:
            raise Exception(f"Error processing DOCX: {str(e)}")
    
    def extract_text_from_excel(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from Excel file"""
        try:
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            text_parts = []
            sheet_info = {}
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                
                # Convert DataFrame to text representation
                sheet_text = f"--- Sheet: {sheet_name} ---\n"
                sheet_text += f"Columns: {', '.join(df.columns.astype(str))}\n\n"
                
                # Add data summary
                sheet_text += f"Data Summary:\n"
                sheet_text += f"Rows: {len(df)}, Columns: {len(df.columns)}\n\n"
                
                # Add actual data (first 100 rows to avoid too much text)
                display_df = df.head(100)
                sheet_text += display_df.to_string(index=False)
                
                if len(df) > 100:
                    sheet_text += f"\n... ({len(df) - 100} more rows)"
                
                text_parts.append(sheet_text)
                sheet_info[sheet_name] = {"rows": len(df), "columns": len(df.columns)}
            
            text = "\n\n".join(text_parts)
            metadata = {
                "sheets": sheet_info,
                "total_sheets": len(excel_file.sheet_names),
                "file_type": "excel"
            }
            
            return text, metadata
            
        except Exception as e:
            raise Exception(f"Error processing Excel: {str(e)}")
    
    def extract_text_from_txt(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            metadata = {
                "lines": len(text.split('\n')),
                "file_type": "txt"
            }
            
            return text, metadata
            
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, 'r', encoding='latin-1') as file:
                text = file.read()
            
            metadata = {
                "lines": len(text.split('\n')),
                "file_type": "txt",
                "encoding": "latin-1"
            }
            
            return text, metadata
            
        except Exception as e:
            raise Exception(f"Error processing TXT: {str(e)}")
    
    def process_document(self, file_path: str, filename: str = None) -> Dict[str, Any]:
        """Process a document and return structured data"""
        if not self.is_supported(file_path):
            raise ValueError(f"Unsupported file type: {Path(file_path).suffix}")
        
        file_ext = Path(file_path).suffix.lower()
        
        # Extract text based on file type
        if file_ext == '.pdf':
            text, metadata = self.extract_text_from_pdf(file_path)
        elif file_ext == '.docx':
            text, metadata = self.extract_text_from_docx(file_path)
        elif file_ext in ['.xlsx', '.xls']:
            text, metadata = self.extract_text_from_excel(file_path)
        elif file_ext == '.txt':
            text, metadata = self.extract_text_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file extension: {file_ext}")
        
        # Generate document hash for deduplication
        doc_hash = hashlib.md5(text.encode()).hexdigest()
        
        # Prepare document data
        document_data = {
            "filename": filename or Path(file_path).name,
            "file_path": file_path,
            "file_type": metadata["file_type"],
            "text": text,
            "metadata": metadata,
            "hash": doc_hash,
            "processed_at": datetime.now().isoformat(),
            "size": len(text)
        }
        
        return document_data
    
    def chunk_text(self, text: str, chunk_size: int = 1000, overlap: int = 200) -> List[Dict[str, Any]]:
        """Split text into overlapping chunks"""
        if len(text) <= chunk_size:
            return [{"text": text, "start": 0, "end": len(text), "chunk_id": 0}]
        
        chunks = []
        start = 0
        chunk_id = 0
        
        while start < len(text):
            end = start + chunk_size
            
            # If we're not at the end, try to break at a sentence or paragraph
            if end < len(text):
                # Look for sentence endings near the end
                for i in range(end, max(start + chunk_size//2, end - 100), -1):
                    if text[i] in '.!?\n':
                        end = i + 1
                        break
            
            chunk_text = text[start:end].strip()
            
            if chunk_text:
                chunks.append({
                    "text": chunk_text,
                    "start": start,
                    "end": end,
                    "chunk_id": chunk_id
                })
                chunk_id += 1
            
            # Move start position with overlap
            start = end - overlap
            
            # Ensure we don't go backwards
            if start <= chunks[-1]["start"] if chunks else False:
                start = chunks[-1]["end"]
        
        return chunks
    
    def process_uploaded_file(self, uploaded_file, chunk_size: int = 1000, overlap: int = 200) -> Dict[str, Any]:
        """Process an uploaded file from Streamlit"""
        # Save uploaded file temporarily
        with tempfile.NamedTemporaryFile(delete=False, suffix=Path(uploaded_file.name).suffix) as temp_file:
            temp_file.write(uploaded_file.getbuffer())
            temp_path = temp_file.name
        
        try:
            # Process the document
            doc_data = self.process_document(temp_path, uploaded_file.name)
            
            # Chunk the text
            chunks = self.chunk_text(doc_data["text"], chunk_size, overlap)
            doc_data["chunks"] = chunks
            doc_data["num_chunks"] = len(chunks)
            
            return doc_data
        
        finally:
            # Clean up temporary file
            os.unlink(temp_path)